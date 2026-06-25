from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "智慧校园外卖点餐系统期末大作业报告-陈天旸.docx"
FIG_DIR = ROOT / "report-assets" / "figures"
SHOT_DIR = ROOT / "report-assets" / "screenshots"
FIG_DIR.mkdir(parents=True, exist_ok=True)

FONT = "Microsoft YaHei"
CN_FONT_PATH = "C:/Windows/Fonts/simhei.ttf"


def font(size=34, bold=False):
    path = "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/Deng.ttf"
    return ImageFont.truetype(path, size)


def set_run_font(run, name=FONT, size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_font(paragraph, name=FONT, size=11, color=None, bold=False):
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size, bold=bold, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7C2D0")


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=9.5, color="555555")
    return p


def add_body(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.28)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for line in dedent(code).strip().splitlines():
        run = p.add_run(line.rstrip() + "\n")
        set_run_font(run, name="Consolas", size=8.5)


def add_table(doc, headers, rows, widths=None, caption=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths:
        for i, w in enumerate(widths):
            table.columns[i].width = Inches(w)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, fill="E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], str(text))
    if caption:
        add_caption(doc, caption)
    return table


def draw_box(draw, xy, text, fill, outline=(46, 116, 181), fs=28, bold=False):
    if isinstance(fill, str) and not fill.startswith("#"):
        fill = "#" + fill
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    lines = str(text).split("\n")
    f = font(fs, bold)
    total_h = sum(draw.textbbox((0, 0), line, font=f)[3] for line in lines) + (len(lines) - 1) * 8
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill=(20, 35, 55), font=f)
        y += bbox[3] - bbox[1] + 8


def arrow(draw, start, end, color=(46, 116, 181)):
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        pts = [(ex, ey), (ex - 14, ey - 9), (ex - 14, ey + 9)]
    else:
        pts = [(ex, ey), (ex + 14, ey - 9), (ex + 14, ey + 9)]
    draw.polygon(pts, fill=color)


def make_diagrams():
    # Use case diagram
    img = Image.new("RGB", (1800, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((640, 45), "智慧校园外卖点餐系统用例图", fill=(11, 37, 69), font=font(44, True))
    actors = {"游客": (120, 250), "客户": (120, 570), "管理员": (1500, 410)}
    for name, (x, y) in actors.items():
        d.ellipse((x, y, x + 90, y + 90), outline=(50, 50, 50), width=4)
        d.line((x + 45, y + 90, x + 45, y + 190), fill=(50, 50, 50), width=4)
        d.line((x - 20, y + 125, x + 110, y + 125), fill=(50, 50, 50), width=4)
        d.line((x + 45, y + 190, x - 10, y + 260), fill=(50, 50, 50), width=4)
        d.line((x + 45, y + 190, x + 100, y + 260), fill=(50, 50, 50), width=4)
        d.text((x + 5, y + 275), name, fill=(0, 0, 0), font=font(30, True))
    cases = [
        ("浏览菜品", 480, 210), ("模糊查询菜名", 480, 350), ("注册/登录", 480, 490),
        ("加入购物车", 760, 350), ("提交订单", 760, 490), ("查询订单进度", 760, 630),
        ("分类管理", 1070, 250), ("菜品管理", 1070, 390), ("订单管理", 1070, 530), ("数据统计/库存预警", 1070, 670),
    ]
    for text, x, y in cases:
        d.ellipse((x, y, x + 230, y + 80), fill=(245, 248, 252), outline=(46, 116, 181), width=3)
        bbox = d.textbbox((0, 0), text, font=font(26))
        d.text((x + (230 - bbox[2]) / 2, y + 23), text, fill=(20, 35, 55), font=font(26))
    for end in [(480, 250), (480, 390)]:
        arrow(d, (250, 340), end)
    for end in [(480, 530), (760, 390), (760, 530), (760, 670)]:
        arrow(d, (250, 660), end)
    for end in [(1300, 290), (1300, 430), (1300, 570), (1300, 710)]:
        arrow(d, (1500, 540), end)
    img.save(FIG_DIR / "fig1-use-case.png")

    # Function module diagram
    img = Image.new("RGB", (1800, 1000), "white")
    d = ImageDraw.Draw(img)
    d.text((610, 40), "系统功能模块图", fill=(11, 37, 69), font=font(46, True))
    draw_box(d, (680, 130, 1120, 230), "智慧校园外卖点餐系统", "E8EEF5", fs=32, bold=True)
    modules = [
        ("游客端\n菜品浏览、分类筛选、模糊搜索", 140, 360, 500, 520),
        ("客户端\n注册登录、购物车、提交订单、订单进度", 570, 360, 1010, 520),
        ("管理员端\n分类、菜品、订单、库存、数据看板", 1080, 360, 1580, 520),
    ]
    for text, x1, y1, x2, y2 in modules:
        draw_box(d, (x1, y1, x2, y2), text, "F4F6F9", fs=26, bold=True)
        arrow(d, (900, 230), ((x1 + x2) // 2, y1))
    details = [
        ("首页展示\n热销推荐\n游客搜索", 150, 650, 480, 810),
        ("购物车管理\n库存校验\n订单状态时间轴", 590, 650, 990, 810),
        ("菜品上下架\n已有订单不删除\n经营统计", 1100, 650, 1560, 810),
    ]
    for text, x1, y1, x2, y2 in details:
        draw_box(d, (x1, y1, x2, y2), text, "FFFFFF", fs=24)
    img.save(FIG_DIR / "fig2-modules.png")

    # Order flow
    img = Image.new("RGB", (1800, 880), "white")
    d = ImageDraw.Draw(img)
    d.text((640, 40), "客户下单业务流程图", fill=(11, 37, 69), font=font(44, True))
    steps = [
        ("浏览/搜索菜品", 120, 230), ("加入购物车", 410, 230), ("填写收货信息", 700, 230),
        ("提交订单", 990, 230), ("库存扣减", 1280, 230), ("待接单", 1280, 500),
        ("管理员接单", 990, 500), ("制作/配送", 700, 500), ("订单完成", 410, 500),
    ]
    for text, x, y in steps:
        draw_box(d, (x, y, x + 220, y + 90), text, "F4F6F9", fs=24, bold=True)
    for i in range(4):
        arrow(d, (steps[i][1] + 220, steps[i][2] + 45), (steps[i + 1][1], steps[i + 1][2] + 45))
    arrow(d, (1390, 320), (1390, 500))
    for i in range(5, 8):
        arrow(d, (steps[i][1], steps[i][2] + 45), (steps[i + 1][1] + 220, steps[i + 1][2] + 45))
    d.text((1140, 375), "库存不足则提示失败，不生成订单", fill=(155, 28, 28), font=font(24, True))
    img.save(FIG_DIR / "fig3-order-flow.png")

    # ER diagram
    img = Image.new("RGB", (1900, 1120), "white")
    d = ImageDraw.Draw(img)
    d.text((730, 35), "数据库 ER 图", fill=(11, 37, 69), font=font(44, True))
    ents = {
        "user\n用户": (90, 160, 410, 310),
        "cart\n购物车": (550, 160, 870, 310),
        "cart_item\n购物车明细": (1010, 160, 1370, 310),
        "category\n菜品分类": (90, 500, 410, 650),
        "dish\n菜品": (550, 500, 870, 650),
        "orders\n订单": (1010, 500, 1370, 650),
        "order_item\n订单明细": (1450, 500, 1810, 650),
        "order_status_log\n订单状态日志": (1010, 820, 1370, 970),
    }
    for text, box in ents.items():
        draw_box(d, box, text, "F4F6F9", fs=24, bold=True)
    rels = [
        ((410,235),(550,235),"1:1"),
        ((870,235),(1010,235),"1:N"),
        ((410,575),(550,575),"1:N"),
        ((870,575),(1010,575),"N:1/下单引用"),
        ((1370,575),(1450,575),"1:N"),
        ((1190,650),(1190,820),"1:N"),
        ((1190,310),(1190,500),"用户下单"),
        ((1190,310),(1190,500),""),
    ]
    arrow(d, (410,235),(550,235)); d.text((455,200),"拥有",font=font(22),fill=(0,0,0))
    arrow(d, (870,235),(1010,235)); d.text((900,200),"包含",font=font(22),fill=(0,0,0))
    arrow(d, (410,575),(550,575)); d.text((435,540),"分类",font=font(22),fill=(0,0,0))
    arrow(d, (870,575),(1010,575)); d.text((900,540),"生成订单",font=font(22),fill=(0,0,0))
    arrow(d, (1370,575),(1450,575)); d.text((1390,540),"包含",font=font(22),fill=(0,0,0))
    arrow(d, (1190,650),(1190,820)); d.text((1210,720),"记录状态",font=font(22),fill=(0,0,0))
    d.line((250,310,250,760,1010,760,1010,575),fill=(46,116,181),width=4)
    d.text((470,730),"用户提交订单",font=font(22),fill=(0,0,0))
    img.save(FIG_DIR / "fig4-er.png")


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("智慧校园外卖点餐系统期末大作业报告")
    set_run_font(r, size=9, color="777777")
    return doc


def cover(doc):
    p = doc.add_paragraph()
    r = p.add_run("成绩")
    set_run_font(r, size=12, bold=True)
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("南京工程学院")
    set_run_font(r, size=22, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("期末大作业课题说明书（论文）")
    set_run_font(r, size=18, bold=True)
    for _ in range(2):
        doc.add_paragraph()
    rows = [
        ("题  目", "智慧校园外卖点餐系统"),
        ("课 程 名 称", "WEB应用开发技术课程设计"),
        ("院（系、部、中心）", "计算机工程学院"),
        ("专       业", "软件工程"),
        ("班       级", "软工中外232"),
        ("学 生 姓 名", "陈天旸"),
        ("学       号", "21230215"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.5)
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], k, bold=True)
        set_cell_text(table.rows[i].cells[1], v)
    set_table_borders(table)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("说明：本报告为完整电子版，正文、图表、系统截图和关键代码分析均已整理。")
    set_run_font(r, size=10, color="666666")
    doc.add_page_break()


def add_toc_note(doc):
    doc.add_heading("报告结构说明", level=1)
    add_body(doc, "本报告依据《WEB开发技术》期末大作业任务书和课程设计报告模板编写，围绕“Java Web外卖点餐系统”选题展开。报告内容覆盖系统设计目的、需求分析、开发技术、功能模块、数据库设计、Web软件系统设计、调试测试、课程收获与参考资料。")
    add_body(doc, "由于本项目实际采用 Spring Boot + MyBatis + React 前后端分离架构，报告在传统 JSP/Servlet/JDBC/Ajax 要求基础上，说明了使用现代 Java Web 框架完成同等业务目标的原因：后端仍由 Java Web 服务提供接口与数据库访问，前端通过 Axios 发起异步请求并渲染页面，能够更清晰地体现数据库动态数据驱动页面的过程。")


def main_content(doc):
    doc.add_heading("一、系统设计目的", level=1)
    add_body(doc, "本系统的设计目的，是围绕校园外卖点餐场景实现一个具有完整业务闭环的 Web 应用。系统面向游客、客户和管理员三类角色：游客可以浏览和搜索菜品，客户登录后可以加入购物车、提交订单并查询订单进度，管理员可以维护菜品分类、菜品信息、订单状态和库存数据。")
    add_body(doc, "在实现过程中，我重点考虑了课程要求中的页面交互、数据库动态数据、Ajax 获取数据和 Java Web 后端业务处理。后端使用 Spring Boot 组织 Controller、Service、Mapper 分层逻辑，MyBatis 负责数据库访问；前端使用 React 构建页面，通过 Axios 与后端接口通信。这样做的原因是前后端职责清晰，接口数据结构明确，既便于调试，也便于在报告中说明每个功能模块的实现过程。")
    add_body(doc, "作为个人独立完成的课程设计，本项目不仅完成基本的浏览、下单和管理功能，还加入了订单状态时间轴、热销推荐、库存预警、数据看板和“已有订单菜品只能下架不能删除”等业务规则，使系统更接近真实外卖点餐平台。")

    doc.add_heading("二、系统需求分析", level=1)
    add_body(doc, "根据任务书要求，本项目选择“Java Web外卖点餐系统”。系统用于餐饮企业或校园食堂实现外卖点餐业务，用户分为管理员、客户、游客。需求分析时，我将功能拆分为前台点餐业务和后台管理业务两部分，并保证所有页面核心数据来源于数据库。")
    add_table(doc, ["角色", "主要需求", "权限边界"], [
        ["游客", "浏览首页、查看菜品分类、模糊查询菜名、查看菜品详情", "不能下单，加入购物车时跳转登录"],
        ["客户", "注册登录、维护个人信息、加入购物车、提交订单、查询订单进度、取消待接单订单", "只能查看和操作自己的订单"],
        ["管理员", "分类管理、菜品管理、图片上传、订单管理、库存预警、销售统计", "可查看全部订单并推动订单状态流转"],
    ], caption="表 1 用户角色与需求分析")
    doc.add_picture(str(FIG_DIR / "fig1-use-case.png"), width=Inches(6.2))
    add_caption(doc, "图 1 系统用例图")
    add_body(doc, "用例图体现了三类用户和系统之间的主要交互。游客侧重点是浏览和搜索；客户侧重点是点餐和订单跟踪；管理员侧重点是后台维护与经营管理。这样划分可以让权限控制更清楚，也避免客户直接访问后台数据。")
    doc.add_picture(str(FIG_DIR / "fig3-order-flow.png"), width=Inches(6.2))
    add_caption(doc, "图 2 客户下单业务流程图")
    add_body(doc, "下单流程中最关键的业务规则是库存校验和状态流转。客户提交订单时系统会根据购物车明细计算总价，并扣减对应菜品库存；订单生成后状态为待接单，管理员按接单、制作、配送、完成的顺序推动订单状态。若客户在待接单阶段取消订单，系统会恢复库存和销量。")

    doc.add_heading("三、开发技术简介", level=1)
    add_body(doc, "本项目采用 Spring Boot、MyBatis、MySQL、React、Vite、Axios、TanStack Query、JWT 等技术完成。选择前后端分离方案，是因为外卖点餐系统有较多列表、筛选、状态刷新和表单交互，前端独立实现页面交互会更灵活；后端通过 RESTful API 输出 JSON 数据，能够清晰体现 Ajax 异步获取数据并渲染页面的过程。")
    add_table(doc, ["技术", "用途", "采用原因"], [
        ["Spring Boot", "后端 Web 服务、接口路由、参数校验", "简化 Java Web 配置，便于快速实现 Controller/Service 分层"],
        ["MyBatis", "数据库访问和 SQL 映射", "SQL 可控，适合展示课程设计中的数据表关系和查询逻辑"],
        ["MySQL", "保存用户、分类、菜品、购物车和订单数据", "关系模型清晰，适合 ER 图和二维关系表设计"],
        ["React + Vite", "前端页面和交互", "组件化开发，页面跳转、状态刷新和表单处理更清楚"],
        ["Axios / React Query", "Ajax 请求、数据缓存和刷新", "页面数据全部从后端接口获取，符合动态渲染要求"],
        ["JWT", "登录状态和接口权限", "前后端分离场景下便于携带身份信息并限制后台接口"],
    ], caption="表 2 开发技术及选择原因")
    add_body(doc, "虽然任务书中列出了 JSP、Servlet、JDBC、JQuery、JSTL 等传统 Java Web 技术点，本项目采用的是更高层的 Spring Boot 和 React 实现。Spring Boot 中的 Controller 可以看作对 Servlet 请求处理的工程化封装，MyBatis 完成 JDBC 数据库访问的封装，React 与 Axios 完成 JQuery/Ajax 类似的异步交互。这样既保留了 Java Web 的核心请求-处理-数据库-响应链路，也提升了系统可维护性。")

    doc.add_heading("四、系统功能模块图", level=1)
    doc.add_picture(str(FIG_DIR / "fig2-modules.png"), width=Inches(6.2))
    add_caption(doc, "图 3 系统功能模块图")
    add_body(doc, "系统模块分为游客端、客户点餐端和管理员后台。游客端提供浏览和搜索入口；客户点餐端完成购物车、订单提交和进度查询；管理员后台完成分类、菜品、订单、库存和统计管理。我的实现重点放在后端业务规则和接口设计，同时通过前端页面将功能完整展示出来。")
    add_bullets(doc, [
        "游客端：降低使用门槛，未登录也可以浏览菜品和分类，体现外卖平台的公开展示属性。",
        "客户端：围绕点餐闭环设计，从加入购物车到订单状态时间轴，让用户知道订单当前进度。",
        "管理员端：围绕运营管理设计，重点是分类维护、菜品上下架、订单处理和库存预警。",
    ])

    doc.add_heading("五、系统的详细设计", level=1)
    doc.add_heading("1、数据库设计", level=2)
    add_body(doc, "数据库设计围绕用户、菜品、购物车、订单四条主线展开。用户表保存账号和角色，分类表与菜品表构成菜品展示基础，购物车表和购物车明细表保存客户临时点餐数据，订单表和订单明细表保存最终交易数据，订单状态日志表用于记录订单每次状态变化。")
    doc.add_picture(str(FIG_DIR / "fig4-er.png"), width=Inches(6.2))
    add_caption(doc, "图 4 数据库 ER 图")
    add_table(doc, ["表名", "主要字段", "设计说明"], [
        ["user", "id、username、password、phone、role、address", "保存客户和管理员账号，role 字段区分权限"],
        ["category", "id、name、sort_order", "保存菜品分类，后台可新增、修改、删除"],
        ["dish", "id、category_id、name、price、stock、image_url、status", "保存菜品信息，status 控制上架/下架"],
        ["cart / cart_item", "user_id、dish_id、quantity", "保存客户购物车和菜品数量"],
        ["orders", "order_no、user_id、total_amount、receiver_*、status、remark", "保存订单主信息、收货信息和备注"],
        ["order_item", "order_id、dish_id、dish_name、dish_price、quantity", "保存订单菜品快照，避免菜品后续修改影响历史订单"],
        ["order_status_log", "order_id、status、operator_id、remark", "记录订单状态变化，用于订单进度时间轴"],
    ], caption="表 3 数据库逻辑关系表")
    add_body(doc, "其中 order_item 保存的是菜品名称和价格快照，这样即使管理员之后修改菜品名称或价格，历史订单仍然保持下单时的真实信息。已有订单引用过的菜品不允许直接删除，而是改为下架，这也是为了保护订单明细的完整性。")

    doc.add_heading("2、WEB软件系统设计", level=2)
    add_body(doc, "系统采用前后端分离结构。前端路由负责页面跳转，后端接口负责业务处理，数据库保存状态。后端接口统一返回 JSON，前端通过 Axios 请求后将数据渲染为页面。下面结合界面截图和关键代码说明主要模块。")
    add_table(doc, ["模块", "主要接口/页面", "实现说明"], [
        ["认证模块", "/api/auth/register、/api/auth/login", "注册区分客户和管理员，管理员需输入预设码 yangyang"],
        ["菜品模块", "/api/dishes、/api/admin/dishes", "支持分类筛选、模糊查询、推荐、上下架、图片上传"],
        ["购物车模块", "/api/cart", "客户登录后维护购物车，结算时从购物车生成订单"],
        ["订单模块", "/api/orders、/api/admin/orders", "客户查自己的订单，管理员查全部订单并修改状态"],
        ["统计模块", "/api/admin/statistics/*", "为后台数据看板、热销排行和库存预警提供数据"],
    ], caption="表 4 主要接口与模块说明")

    doc.add_heading("（1）首页、菜品浏览与模糊查询", level=3)
    for name, cap in [
        ("01-home.png", "图 5 系统首页：展示菜品分类、今日推荐和热销榜"),
        ("02-dish-list.png", "图 6 菜品列表：支持关键词、分类和排序筛选"),
        ("03-dish-detail.png", "图 7 菜品详情：显示图片、价格、库存、分类和加入购物车入口"),
    ]:
        doc.add_picture(str(SHOT_DIR / name), width=Inches(6.2))
        add_caption(doc, cap)
    add_body(doc, "首页和菜品页的数据均来自后端数据库。前端通过查询参数传递 keyword、categoryId、sort 等条件，后端 MyBatis 根据条件动态拼接查询语句。这样做的原因是筛选逻辑集中在后端，前端只负责展示和交互，后续扩展价格排序、销量排序或分页时更方便。")
    add_code(doc, """
@GetMapping
public ApiResponse<PageResult<DishVO>> list(DishQuery query) {
    return ApiResponse.success(dishService.list(query));
}

@Select(\"\"\"
SELECT d.*, c.name AS category_name
FROM dish d LEFT JOIN category c ON c.id = d.category_id
<where>
  <if test="keyword != null and keyword != ''">
    AND d.name LIKE CONCAT('%', #{keyword}, '%')
  </if>
  <if test="categoryId != null">AND d.category_id = #{categoryId}</if>
</where>
\"\"\")
List<Dish> findPage(DishQuery query);
""")

    doc.add_heading("（2）注册登录与权限控制", level=3)
    for name, cap in [
        ("04-register-customer.png", "图 8 客户注册界面"),
        ("05-register-admin-code.png", "图 9 管理员注册界面：选择管理员后显示预设码输入框"),
        ("06-login.png", "图 10 登录界面"),
    ]:
        doc.add_picture(str(SHOT_DIR / name), width=Inches(6.2))
        add_caption(doc, cap)
    add_body(doc, "注册功能按照任务书要求区分管理员和客户。普通客户可以直接注册；管理员注册必须输入预设码 yangyang。这个校验放在后端而不是只放在前端，是为了防止用户绕过页面直接调用接口创建管理员账号。")
    add_code(doc, """
private String resolveRegisterRole(RegisterRequest request) {
    String role = StringUtils.hasText(request.getRole())
        ? request.getRole().trim().toUpperCase()
        : ROLE_CUSTOMER;
    if (ROLE_CUSTOMER.equals(role)) return ROLE_CUSTOMER;
    if (ROLE_ADMIN.equals(role)) {
        if (!ADMIN_REGISTER_CODE.equals(request.getAdminCode())) {
            throw new BusinessException(400, "管理员注册码错误");
        }
        return ROLE_ADMIN;
    }
    throw new BusinessException(400, "注册角色无效");
}
""")
    add_body(doc, "登录成功后后端生成 JWT，前端保存 token 并在后续请求头中携带。后台接口会通过 CurrentUserUtil.requireAdmin 校验角色，避免普通客户访问管理端接口。")

    doc.add_heading("（3）购物车、结算与订单进度", level=3)
    for name, cap in [
        ("07-cart.png", "图 11 购物车界面：客户确认菜品数量并进入结算"),
        ("08-checkout.png", "图 12 结算界面：填写收货人、电话、地址和备注"),
        ("09-my-orders.png", "图 13 我的订单：客户查看自己的订单列表"),
        ("10-order-detail-progress.png", "图 14 订单详情：展示订单进度时间轴和菜品明细"),
    ]:
        doc.add_picture(str(SHOT_DIR / name), width=Inches(6.2))
        add_caption(doc, cap)
    add_body(doc, "订单模块是系统的核心业务。客户提交订单时，系统从购物车读取菜品明细，校验菜品是否上架和库存是否足够，计算总金额，写入订单主表和订单明细表，并记录第一条订单状态日志。")
    add_code(doc, """
@Transactional
public CreateOrderVO create(Long userId, CreateOrderRequest request) {
    List<CartItem> cartItems = cartMapper.findItemsByUserId(userId);
    validateCartItems(cartItems);
    Order order = new Order();
    order.setOrderNo(generateOrderNo());
    order.setUserId(userId);
    order.setTotalAmount(totalAmount(cartItems));
    orderMapper.insertOrder(order);
    addStatusLog(order.getId(), "PENDING", userId, "CUSTOMER", "提交订单");
    for (CartItem item : cartItems) {
        dishMapper.decreaseStock(item.getDishId(), item.getQuantity());
        dishMapper.increaseSales(item.getDishId(), item.getQuantity());
        orderMapper.insertOrderItem(toOrderItem(order.getId(), item));
    }
    cartMapper.clearByUserId(userId);
    return toCreateOrderVO(order);
}
""")
    add_body(doc, "订单进度使用 order_status_log 表记录。这样做的好处是订单当前状态和历史状态分离，前端可以根据日志生成时间轴，而不是只显示一个状态文本。")

    doc.add_heading("（4）管理员后台管理", level=3)
    for name, cap in [
        ("11-admin-dashboard.png", "图 15 管理端数据看板：展示订单、销售额、客户数、热销菜品等统计"),
        ("12-admin-categories.png", "图 16 分类管理：新增、修改、删除菜品分类"),
        ("13-admin-dishes.png", "图 17 菜品管理：维护菜品分类、价格、库存、图片和上下架状态"),
        ("14-admin-orders.png", "图 18 订单管理：查看所有客户订单并推动状态流转"),
        ("15-admin-stock.png", "图 19 库存预警：提示库存不足或需要补货的菜品"),
    ]:
        doc.add_picture(str(SHOT_DIR / name), width=Inches(6.2))
        add_caption(doc, cap)
    add_body(doc, "后台模块用于管理员维护经营数据。分类管理中，如果分类下已有菜品，后端会阻止删除；菜品管理中，如果菜品已经出现在订单明细中，则删除操作会转为下架。这两个规则都是为了保护数据库关联数据，避免历史订单和菜品分类出现断裂。")
    add_code(doc, """
@Transactional
public boolean delete(Long id) {
    if (dishMapper.findById(id) == null) {
        throw new BusinessException(404, "菜品不存在");
    }
    if (dishMapper.countOrderItems(id) > 0) {
        dishMapper.updateStatus(id, 0);
        return false;
    }
    dishMapper.delete(id);
    return true;
}
""")
    add_body(doc, "管理员订单管理支持按订单号、用户名、手机号、状态和日期查询订单。订单状态必须按流程流转，不能跳过中间状态，这样可以模拟真实外卖业务中接单、制作、配送、完成的过程。")

    doc.add_heading("（5）调试过程中的问题及解决方法", level=3)
    add_table(doc, ["问题", "原因分析", "解决方法"], [
        ["前端跨域请求失败", "前端运行在 5173，后端运行在 8090，浏览器会触发跨域限制", "后端配置 CORS，允许本地前端访问接口"],
        ["登录后后台接口仍提示无权限", "请求头未正确携带 JWT 或用户角色不是 ADMIN", "在 Axios 拦截器中统一加入 Authorization，并在后端 requireAdmin 校验角色"],
        ["订单提交后库存不同步", "只写入订单但没有扣减 dish.stock", "下单事务中同时扣减库存、增加销量、清空购物车"],
        ["已有订单菜品被删除后历史订单不完整", "订单明细引用 dish_id，直接删除会影响数据关联", "如果菜品已有订单，删除改为下架"],
        ["截图中订单页面为空", "演示数据库没有客户订单", "通过演示客户提交订单后再截图，保证报告展示完整业务闭环"],
    ], caption="表 5 调试问题及解决方法")

    doc.add_heading("（6）程序的不足与改进", level=3)
    add_body(doc, "当前系统已经完成课程要求的核心业务，但仍有改进空间。首先，系统还没有接入真实支付能力，订单金额只是用于业务展示；其次，评价功能、配送员端和商家多门店管理还可以继续扩展；第三，管理员权限目前只有 ADMIN 和 CUSTOMER 两类，后续可以拆分为菜品管理员、订单管理员和统计查看员；最后，项目目前主要面向桌面端展示，移动端适配还可以进一步优化。")
    add_body(doc, "如果继续完善，我会优先增加评价模块、支付状态、配送员端和更细粒度的权限控制，并将系统部署到服务器，配置生产环境数据库和文件上传路径，使项目从课程设计进一步接近真实可用系统。")

    doc.add_heading("六、课程设计收获体会", level=1)
    add_body(doc, "通过本次《WEB开发技术》期末大作业，我完整经历了一个 Web 项目从需求分析、数据库设计、接口设计、前端页面实现、后端业务实现、联调测试到报告整理的过程。相比只写单个页面或单个接口，一个完整外卖点餐系统更强调模块之间的数据流转。例如客户从菜品列表加入购物车，购物车生成订单，订单扣减库存，管理员修改订单状态，客户再查看订单进度，这些步骤必须通过数据库和接口串联起来，任何一个环节设计不清楚都会影响整体体验。")
    add_body(doc, "在后端实现中，我更深刻地理解了分层设计的意义。Controller 只负责接收请求和返回响应，Service 负责业务规则，Mapper 负责数据库访问。这样写虽然比把所有逻辑堆在一个类里更繁琐，但当业务变复杂时优势非常明显。例如“已有订单的菜品只能下架不能删除”这个需求，只需要在 DishService 的删除逻辑中增加订单明细数量判断，就可以同时保护接口和前端操作。")
    add_body(doc, "在数据库设计方面，我认识到订单类系统不能只保存当前菜品信息，还要保存订单当时的菜品快照。如果订单明细只依赖 dish 表，一旦管理员修改菜品名称或价格，历史订单就会失真。因此我在 order_item 表中保存 dish_name、dish_price、quantity、subtotal 等字段，保证订单历史记录稳定。这种设计让我理解了数据库不仅要满足当前页面展示，还要考虑未来数据追溯和业务安全。")
    add_body(doc, "在前端实现中，我体会到用户体验不仅是页面好看，还包括流程是否清楚、反馈是否及时。例如未登录用户加入购物车时跳转登录，订单详情中使用时间轴展示进度，管理员后台用状态标签区分上架、下架、待接单和已完成等状态。这些细节能让系统更容易理解，也更适合验收展示。")
    add_body(doc, "本次项目也让我发现自己在项目规划和调试上的不足。最开始如果只关注页面效果，很容易忽略权限校验、异常提示和数据一致性；后面补充管理员注册码、菜品删除保护和订单状态日志后，系统的完整性明显提升。以后再做 Web 项目时，我会更早地从业务规则、数据模型和异常场景出发，再设计页面和接口。总体来说，本次课程设计让我把 HTML、CSS、JavaScript、Ajax、Java 后端和数据库知识串成了一个完整实践，对 Web 应用开发流程有了更系统的认识。")

    doc.add_heading("七、主要参考资料", level=1)
    refs = [
        "[1] 张娜. Java Web开发技术教程（第三版）[M]. 北京：清华大学出版社，2023.",
        "[2] 黑马程序员. JavaEE企业级应用开发教程（Spring+SpringMVC+MyBatis）[M]. 北京：人民邮电出版社，2017.",
        "[3] 孙卫琴. Tomcat与Java Web开发技术详解（第三版）[M]. 北京：电子工业出版社，2019.",
        "[4] Spring Boot Reference Documentation. https://docs.spring.io/spring-boot/.",
        "[5] MyBatis 参考文档. https://mybatis.org/mybatis-3/.",
        "[6] React Documentation. https://react.dev/.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        set_run_font(r, size=10.5)


def evaluation_page(doc):
    doc.add_page_break()
    doc.add_heading("期末大作业成绩评价", level=1)
    rows = [
        ("动手能力（40%）", ""),
        ("创新能力（10%）", ""),
        ("大作业总结报告（20%）", ""),
        ("总评成绩", ""),
        ("大作业综合成绩评定结论", ""),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(2.4)
    table.columns[1].width = Inches(4.0)
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], k, bold=True, fill="F2F4F7")
        set_cell_text(table.rows[i].cells[1], v)
        table.rows[i].height = Inches(0.55)
    set_table_borders(table)
    add_body(doc, "以上评价栏保留给任课教师填写。", first_line=False)


def build():
    make_diagrams()
    doc = setup_doc()
    cover(doc)
    add_toc_note(doc)
    main_content(doc)
    evaluation_page(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
